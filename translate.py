#!/usr/bin/env python
# -*- coding: utf-8 -*-
# author;cuich time:2019/09/29
from googletrans import Translator
from docx import Document
import time
import os
# 进行判断决定是否新建docx文件
time = time.strftime("%Y-%m-%d %H:%M", time.localtime())
filename = r'D:\trans.docx'
if os.path.exists(filename):
    print('')
else:
    document = Document()
    document.add_heading('认真学习是一件最有趣的事了', level=0)
    document.save(filename)

# 翻译加保存
newdocument = Document(filename)
newdocument.add_paragraph(time)
choose = input("英译中：1"+'\n'+'中译英：2'+'\n'+'请选择1或2： '+'使用过程中按0退出： ')
if choose == '1':
    while True:
        words = input("请输入要翻译的单词或句子： ")
        if words == '0':
            print('祝考试成功！')
            break
        else:
            translator = Translator()
            result = translator.translate(words, dest='zh-CN').text
            print(result)
            newdocument.add_paragraph(words + ':'+result)
            newdocument.save(filename)
# 中译英
else:
    while True:
        words = input("请输入要翻译的单词或句子： ")
        if words == '0':
            print('祝考试成功！')
            break
        else:
            translator = Translator()
            result = translator.translate(words).text
            print(result)
            newdocument.add_paragraph(words + ':' + result)
            newdocument.save(filename)
